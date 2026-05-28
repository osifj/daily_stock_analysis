# -*- coding: utf-8 -*-
"""Tests for Markdown-to-Word report conversion used by WeChat file delivery."""
from pathlib import Path

from docx import Document

from scripts.send_wechat_report_file import build_combined_markdown, find_report_files, markdown_to_docx


def test_markdown_to_docx_preserves_headings_table_and_bold(tmp_path: Path) -> None:
    output = tmp_path / "report.docx"
    markdown = """# 个股分析

## 华海清科

**结论**：观望。

| 指标 | 数值 |
|------|------|
| MA5 | 100 |
| 评分 | 70 |

- 风险：波动较大
1. 检查支撑位
"""

    markdown_to_docx(markdown, output, title="测试报告")

    assert output.exists()
    document = Document(output)
    texts = [paragraph.text for paragraph in document.paragraphs]
    assert "测试报告" in texts
    assert "个股分析" in texts
    assert "华海清科" in texts
    assert any("结论" in text and "观望" in text for text in texts)
    assert document.tables
    table = document.tables[0]
    assert table.cell(0, 0).text == "指标"
    assert table.cell(1, 0).text == "MA5"
    assert table.cell(2, 1).text == "70"


def test_find_report_files_and_build_combined_markdown(tmp_path: Path) -> None:
    (tmp_path / "ignored.txt").write_text("ignore", encoding="utf-8")
    stock_report = tmp_path / "report_20260528.md"
    market_report = tmp_path / "market_review_20260528.md"
    stock_report.write_text("# stock", encoding="utf-8")
    market_report.write_text("# market", encoding="utf-8")

    files = find_report_files(tmp_path)
    assert stock_report in files
    assert market_report in files

    combined = build_combined_markdown(files)
    assert "# 个股分析" in combined
    assert "# 大盘复盘" in combined
    assert "# stock" in combined
    assert "# market" in combined
