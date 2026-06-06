# -*- coding: utf-8 -*-
"""Convert generated Markdown reports to Word and send them to WeChat Work.

Enhanced version with:
- Cover page
- Auto TOC
- Color-coded up/down (red up, green down for Chinese markets)
- Page breaks between stock sections
- Code block gray background
- Header & footer with page numbers
"""
from __future__ import annotations

import argparse
import logging
import re
import sys
from datetime import datetime, timezone, timedelta
from pathlib import Path
from typing import Iterable, List, Optional, Sequence

REPO_ROOT = Path(__file__).resolve().parents[1]
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Cm

from src.config import get_config
from src.notification_sender.wechat_sender import WechatSender


logger = logging.getLogger(__name__)
CN_TZ = timezone(timedelta(hours=8))

# ── Color constants (Chinese market: red = up, green = down) ──
RED_UP = RGBColor(0xCC, 0x00, 0x00)      # 涨红
GREEN_DOWN = RGBColor(0x00, 0x88, 0x00)  # 跌绿
TABLE_HEADER_BG = "D9EAF7"
CODE_BG = "F0F0F0"
BORDER_COLOR = "B7C0CC"
COVER_ACCENT = "1A5276"

# ── Pattern for color detection in table cells ──
_UP_DOWN_RE = re.compile(r'[+-]\d+\.?\d*%?')
_UP_KEYWORDS = re.compile(r'[涨↑📈🚀]|看多|买入|增持|突破|利好')
_DOWN_KEYWORDS = re.compile(r'[跌↓📉🔻]|看空|卖出|减持|跌破|利空')


def _detect_cell_color(text: str):
    """Return (RED_UP | GREEN_DOWN | None) for a table cell's content."""
    t = text.strip()
    if not t:
        return None
    # Explicit + / - numbers
    if t.startswith('+'):
        return RED_UP
    if t.startswith('-'):
        return GREEN_DOWN
    # Chinese keywords
    if _UP_KEYWORDS.search(t):
        return RED_UP
    if _DOWN_KEYWORDS.search(t):
        return GREEN_DOWN
    # Percentage pattern anywhere
    m = _UP_DOWN_RE.search(t)
    if m:
        val = m.group(0)
        if val.startswith('+'):
            return RED_UP
        if val.startswith('-'):
            return GREEN_DOWN
    return None


# ── Low-level XML helpers ──────────────────────────────────────────

def _set_cell_shading(cell, fill: str = TABLE_HEADER_BG) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()  # noqa: SLF001
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _set_cell_border(cell) -> None:
    tc = cell._tc  # noqa: SLF001
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER_COLOR)


def _set_paragraph_shading(paragraph, fill: str) -> None:
    """Add background shading to a paragraph (used for code blocks)."""
    pPr = paragraph._p.get_or_add_pPr()  # noqa: SLF001
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)


def _insert_page_number(footer) -> None:
    """Insert '第 X 页' field into footer."""
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run_with_font(para, "第 ", "Microsoft YaHei", Pt(9))

    # PAGE field
    run = para.add_run()
    run.font.size = Pt(9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)  # noqa: SLF001

    run2 = para.add_run()
    run2.font.size = Pt(9)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._r.append(instr)  # noqa: SLF001

    run3 = para.add_run()
    run3.font.size = Pt(9)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run3._r.append(fld_end)  # noqa: SLF001

    _add_run_with_font(para, " 页", "Microsoft YaHei", Pt(9))


def _add_run_with_font(paragraph, text: str, font_name: str, size: Pt, bold: bool = False, color=None) -> None:
    run = paragraph.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)  # noqa: SLF001
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _add_field(paragraph, field_code: str, display_text: str = "") -> None:
    """Insert a Word field (TOC, etc.)."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)  # noqa: SLF001

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    run2._r.append(instr)  # noqa: SLF001

    if display_text:
        run3 = paragraph.add_run()
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        run3._r.append(fld_sep)  # noqa: SLF001
        run4 = paragraph.add_run(display_text)

    run5 = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run5._r.append(fld_end)  # noqa: SLF001


# ── Document setup ──────────────────────────────────────────────────

def _configure_document(document: DocumentObject) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Microsoft YaHei"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")  # noqa: SLF001
    normal_style.font.size = Pt(10.5)
    normal_style.paragraph_format.space_after = Pt(4)
    normal_style.paragraph_format.line_spacing = 1.25

    for style_name, size, space_before, space_after in (
        ("Heading 1", 18, Pt(18), Pt(8)),
        ("Heading 2", 15, Pt(14), Pt(6)),
        ("Heading 3", 13, Pt(10), Pt(4)),
    ):
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")  # noqa: SLF001
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
        style.paragraph_format.space_before = space_before
        style.paragraph_format.space_after = space_after


# ── Cover page ──────────────────────────────────────────────────────

def _add_cover_page(document: DocumentObject, report_count: int) -> None:
    """Add a professional cover page."""
    now = datetime.now(CN_TZ)

    # Empty space at top
    for _ in range(6):
        document.add_paragraph()

    # Title
    title_para = document.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run_with_font(title_para, "股票智能分析报告", "Microsoft YaHei", Pt(28), bold=True,
                       color=RGBColor(0x1A, 0x52, 0x76))

    # Subtitle line
    document.add_paragraph()
    sub_para = document.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run_with_font(sub_para, "Daily Stock Analysis Report", "Consolas", Pt(12),
                       color=RGBColor(0x7F, 0x8C, 0x8D))

    # Decorative line
    document.add_paragraph()
    line_para = document.add_paragraph()
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run_with_font(line_para, "━" * 40, "Microsoft YaHei", Pt(10),
                       color=RGBColor(0xBD, 0xC3, 0xC7))

    # Metadata
    document.add_paragraph()
    meta_items = [
        f"📅 生成日期：{now.strftime('%Y年%m月%d日')}",
        f"⏰ 生成时间：{now.strftime('%H:%M:%S')}（北京时间）",
        f"📊 分析报告：{report_count} 份",
        f"🤖 AI 模型：DeepSeek",
    ]
    for item in meta_items:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run_with_font(p, item, "Microsoft YaHei", Pt(11), color=RGBColor(0x34, 0x49, 0x5E))

    # End cover page
    document.add_page_break()


# ── Table of Contents ───────────────────────────────────────────────

def _add_toc(document: DocumentObject) -> None:
    """Insert auto TOC field."""
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_run_with_font(heading, "📑 目  录", "Microsoft YaHei", Pt(18), bold=True,
                       color=RGBColor(0x1A, 0x52, 0x76))

    document.add_paragraph()  # spacer

    toc_para = document.add_paragraph()
    _add_field(toc_para, ' TOC \\o "1-3" \\h \\z \\u ',
               display_text="（请在 Word 中右键点击此处 → 更新域 以生成目录）")

    document.add_page_break()


# ── Header & Footer ─────────────────────────────────────────────────

def _add_header_footer(document: DocumentObject) -> None:
    """Add header with report title and footer with page numbers."""
    section = document.sections[0]

    # Header
    header = section.header
    header.is_linked_to_previous = False
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run_with_font(header_para, "股票智能分析报告  |  ",
                       "Microsoft YaHei", Pt(8), color=RGBColor(0x95, 0xA5, 0xA6))
    _add_run_with_font(header_para, datetime.now(CN_TZ).strftime('%Y-%m-%d'),
                       "Consolas", Pt(8), color=RGBColor(0x95, 0xA5, 0xA6))

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    _insert_page_number(footer)


# ── Inline Markdown rendering ───────────────────────────────────────

_INLINE_MD_RE = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')


def _add_inline_markdown(paragraph, text: str) -> None:
    """Render inline bold, italic, and code spans."""
    position = 0
    for match in _INLINE_MD_RE.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")  # noqa: SLF001
            run.font.color.rgb = RGBColor(0x8E, 0x44, 0xAD)
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


# ── Code blocks ─────────────────────────────────────────────────────

def _add_code_block(document: DocumentObject, code_text: str) -> None:
    """Add a code block with gray background."""
    paragraph = document.add_paragraph(style="No Spacing")
    _set_paragraph_shading(paragraph, CODE_BG)
    paragraph.paragraph_format.left_indent = Cm(0.5)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(code_text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")  # noqa: SLF001
    run.font.size = Pt(9)


# ── Table parsing & rendering ───────────────────────────────────────

def _is_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def _parse_table(lines: Sequence[str], start: int) -> tuple[Optional[List[List[str]]], int]:
    if start + 1 >= len(lines):
        return None, start
    if "|" not in lines[start] or not _is_table_separator(lines[start + 1]):
        return None, start

    rows: List[List[str]] = []
    index = start
    while index < len(lines) and "|" in lines[index] and lines[index].strip():
        if not _is_table_separator(lines[index]):
            rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
        index += 1
    return rows, index


def _add_table(document: DocumentObject, rows: Sequence[Sequence[str]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    table.autofit = True

    for row_index, row in enumerate(rows):
        for col_index in range(column_count):
            cell = table.cell(row_index, col_index)
            cell_text = row[col_index] if col_index < len(row) else ""
            _set_cell_border(cell)

            # Clear default paragraph and add styled text
            cell.paragraphs[0].clear()
            cell_para = cell.paragraphs[0]
            cell_para.paragraph_format.space_before = Pt(2)
            cell_para.paragraph_format.space_after = Pt(2)

            if row_index == 0:
                # Header row
                _set_cell_shading(cell)
                _add_run_with_font(cell_para, cell_text, "Microsoft YaHei", Pt(9.5),
                                   bold=True, color=RGBColor(0x1A, 0x52, 0x76))
            else:
                # Data row — apply color coding
                color = _detect_cell_color(cell_text)
                _add_run_with_font(cell_para, cell_text, "Microsoft YaHei", Pt(9.5),
                                   color=color)


# ── Main conversion ─────────────────────────────────────────────────

def markdown_to_docx(markdown: str, output_path: Path, *, title: str, report_count: int = 1) -> Path:
    document = Document()
    _configure_document(document)

    # ── Cover page ──
    _add_cover_page(document, report_count)

    # ── TOC ──
    _add_toc(document)

    # ── Header & Footer ──
    _add_header_footer(document)

    # ── Content ──
    lines = markdown.splitlines()
    index = 0
    first_h1_seen = False
    in_code_block = False
    code_lines: List[str] = []

    while index < len(lines):
        raw_line = lines[index].rstrip()
        line = raw_line.strip()

        # Code block start/end
        if line.startswith("```"):
            if in_code_block:
                _add_code_block(document, "\n".join(code_lines))
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            index += 1
            continue

        if in_code_block:
            code_lines.append(raw_line)
            index += 1
            continue

        # Tables
        table_rows, next_index = _parse_table(lines, index)
        if table_rows is not None:
            _add_table(document, table_rows)
            index = next_index
            continue

        # Empty lines / separators
        if not line:
            index += 1
            continue
        if re.fullmatch(r"[-*_]{3,}", line):
            document.add_paragraph()  # visual spacer
            index += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            heading_text = heading_match.group(2).strip()
            if level == 1:
                if first_h1_seen:
                    document.add_page_break()
                else:
                    first_h1_seen = True
            document.add_heading(heading_text, level=level)
            index += 1
            continue

        # Bullet lists
        bullet_match = re.match(r"^[-*+]\s+(.+)$", line)
        if bullet_match:
            paragraph = document.add_paragraph(style="List Bullet")
            _add_inline_markdown(paragraph, bullet_match.group(1).strip())
            index += 1
            continue

        # Numbered lists
        numbered_match = re.match(r"^\d+[.)]\s+(.+)$", line)
        if numbered_match:
            paragraph = document.add_paragraph(style="List Number")
            _add_inline_markdown(paragraph, numbered_match.group(1).strip())
            index += 1
            continue

        # Regular paragraph
        paragraph = document.add_paragraph()
        _add_inline_markdown(paragraph, line)
        index += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


# ── File discovery & combination ────────────────────────────────────

def find_report_files(reports_dir: Path) -> List[Path]:
    patterns = ("report_*.md", "market_review_*.md")
    files: List[Path] = []
    for pattern in patterns:
        files.extend(reports_dir.glob(pattern))
    return sorted(files, key=lambda path: path.stat().st_mtime)


def build_combined_markdown(report_files: Iterable[Path]) -> str:
    parts: List[str] = []
    for report_file in report_files:
        title = "大盘复盘" if report_file.name.startswith("market_review_") else "个股分析"
        parts.append(f"# {title}\n\n" + report_file.read_text(encoding="utf-8"))
    return "\n\n---\n\n".join(parts)


# ── Send ────────────────────────────────────────────────────────────

def send_wechat_report_file(reports_dir: Path, output_path: Optional[Path] = None) -> Path:
    report_files = find_report_files(reports_dir)
    if not report_files:
        raise FileNotFoundError(f"未找到可转换的 Markdown 报告: {reports_dir}")

    output = output_path or reports_dir / f"stock_analysis_report_{datetime.now(CN_TZ).strftime('%Y%m%d_%H%M%S')}.docx"
    markdown = build_combined_markdown(report_files)
    markdown_to_docx(markdown, output, title="股票智能分析报告", report_count=len(report_files))

    sender = WechatSender(get_config())
    summary = (
        "📊 股票智能分析报告已生成，完整内容请查看 Word 附件。\n"
        f"📋 包含报告：{', '.join(path.name for path in report_files)}"
    )
    if not sender.send_to_wechat(summary):
        raise RuntimeError("企业微信报告摘要发送失败")
    if not sender.send_file_to_wechat(output):
        raise RuntimeError("企业微信 Word 报告文件发送失败")
    return output


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Send generated reports as a WeChat Work Word file")
    parser.add_argument("--reports-dir", default="reports", help="Report directory containing Markdown reports")
    parser.add_argument("--output", default="", help="Optional .docx output path")
    return parser.parse_args()


def main() -> int:
    logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
    args = parse_args()
    output = Path(args.output) if args.output else None
    try:
        docx_path = send_wechat_report_file(Path(args.reports_dir), output)
        logger.info("企业微信 Word 报告已发送: %s", docx_path)
        return 0
    except Exception as exc:  # noqa: BLE001
        logger.exception("企业微信 Word 报告发送失败: %s", exc)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
